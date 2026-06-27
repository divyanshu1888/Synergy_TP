from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16) if level == 1 else Pt(13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_page_break():
    doc.add_page_break()

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(
        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    )

# TITLE PAGE
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Software / ML Domain Taskphase")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run("Conceptual Report: Tasks 1 to 5")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 0, 0)

fields = [
    ("Participant Name", "Penta Divyanshu"),
    ("Roll Number / ID", "251090051426"),
    ("Domain", "Software / ML"),
    ("Repository", "https://github.com/divyanshu1888/Synergy_TP"),
    ("Submission Date", "June 27, 2026"),
]
for label, value in fields:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + ":  ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0, 0, 0)

add_page_break()

# ABSTRACT
add_heading("Abstract")
add_body(
    "This report presents the theoretical foundations and conceptual understanding developed "
    "through the Software/ML Domain Taskphase. It covers five tasks spanning version control "
    "with Git and GitHub, Python programming fundamentals, Linux command-line usage, CSV data "
    "parsing, data cleaning with pandas, and data visualization using matplotlib. Each section "
    "explains the purpose of the tools and methods used, the reasoning behind design decisions, "
    "and the significance of each stage in the data processing pipeline. The report focuses on "
    "concepts and analytical thinking rather than code walkthroughs, aiming to demonstrate a "
    "clear understanding of how raw data is transformed into meaningful visual insights through "
    "systematic, well-documented steps."
)
add_page_break()

# 1. INTRODUCTION
add_heading("1. Introduction")
add_body(
    "The Synergy Taskphase is a structured learning program designed to build practical skills "
    "in software development and machine learning workflows. Participants complete a series of "
    "tasks that progressively introduce core concepts: from setting up a development environment "
    "and writing Python scripts, to cleaning real-world datasets and visualizing results."
)
add_body(
    "This report covers the conceptual background for Tasks 1 through 5. Task 1 focused on "
    "establishing a proper development environment using Git, GitHub, virtual environments, and "
    "Linux tools. Task 2 reinforced Python fundamentals through practical scripting exercises. "
    "Task 3 introduced CSV parsing both manually and using pandas, building an understanding of "
    "how structured data is read and summarized. Task 4 extended this by tackling messy, "
    "real-world data that required systematic cleaning before any analysis. Task 5 completed the "
    "pipeline by generating meaningful visualizations from the cleaned data."
)
add_body(
    "The goal of this report is to explain the reasoning behind each tool and technique chosen, "
    "not to describe the code line by line. Understanding why a particular approach is taken is "
    "as important as knowing how to implement it."
)
add_page_break()

# 2. DEVELOPMENT ENVIRONMENT AND VERSION CONTROL
add_heading("2. Development Environment and Version Control")
add_heading("2.1 Git and GitHub", level=2)
add_body(
    "Git is a distributed version control system that tracks changes to files over time. Every "
    "change to the codebase is recorded as a commit, which includes a message, a timestamp, and "
    "a reference to the previous state. This makes it possible to review the history of a "
    "project, revert to earlier versions, and understand when and why specific changes were made."
)
add_body(
    "GitHub is a cloud-based hosting platform for Git repositories. It enables collaboration, "
    "visibility, and structured submission. For this taskphase, GitHub serves as the central "
    "submission point: each task is pushed to the Synergy_TP repository in its own folder, and "
    "the final commit hash is submitted as proof of completion."
)
add_body(
    "Clean version control means writing meaningful commit messages, not committing unnecessary "
    "files such as virtual environment folders or compiled bytecode, and maintaining a logical "
    "folder structure. A repository that is easy to navigate and review reflects professional "
    "development habits."
)

add_heading("2.2 Repository Structure", level=2)
add_body(
    "The Synergy_TP repository follows a flat task-based structure where each task occupies its "
    "own folder at the root level. This separation ensures that tasks do not interfere with each "
    "other and that any reviewer can navigate directly to the relevant folder. Each task folder "
    "contains its own README, source code, data, and output directories, making the repository "
    "self-contained and reproducible."
)

add_heading("2.3 Virtual Environments and requirements.txt", level=2)
add_body(
    "A Python virtual environment is an isolated directory containing a specific Python "
    "interpreter and a set of installed packages. The purpose of isolation is to prevent "
    "dependency conflicts between projects. When two projects require different versions of the "
    "same library, virtual environments ensure they do not interfere with each other."
)
add_body(
    "The requirements.txt file lists all packages and their versions that are needed to run the "
    "project. Any collaborator or reviewer can recreate the exact environment by running "
    "pip install -r requirements.txt, ensuring consistent and reproducible results across machines."
)

add_heading("2.4 The .gitignore File", level=2)
add_body(
    "The .gitignore file tells Git which files and folders to exclude from version tracking. "
    "The virtual environment folder, Python bytecode files with the .pyc extension, the "
    "__pycache__ directory, and operating system files such as .DS_Store on macOS are typically "
    "excluded. Committing these would add unnecessary noise to the repository and could cause "
    "issues when the project is run on a different machine."
)
add_page_break()

# 3. PYTHON AND FILE HANDLING CONCEPTS
add_heading("3. Python and File Handling Concepts")
add_heading("3.1 Functions and Modularity", level=2)
add_body(
    "Functions are the primary unit of reusable logic in Python. Breaking a program into "
    "well-named functions makes the code easier to read, test, and maintain. Each function "
    "should have a single clear responsibility. For the taskphase assignments, each required "
    "function was designed to handle one stage of the pipeline: reading data, converting types, "
    "calculating summaries, or writing output."
)

add_heading("3.2 Lists and Dictionaries", level=2)
add_body(
    "Lists store ordered sequences of items and are used when the position of elements matters "
    "or when items need to be iterated. Dictionaries store key-value pairs and are used when "
    "data needs to be accessed by name rather than by index. In the CSV parsing task, each row "
    "of the CSV was represented as a dictionary where the keys were column headers and the "
    "values were the corresponding row values. This made the data human-readable and easy to "
    "process by field name."
)

add_heading("3.3 File I/O and Context Managers", level=2)
add_body(
    "Reading from and writing to files is done using the built-in open() function. The preferred "
    "pattern is to use a context manager with the 'with' keyword, which automatically closes the "
    "file when the block exits, even if an error occurs. This prevents resource leaks and is "
    "considered the correct approach for file handling in Python."
)

add_heading("3.4 Exception Handling", level=2)
add_body(
    "Exception handling using try and except blocks prevents a program from crashing when it "
    "encounters unexpected input. In the manual CSV parser, exception handling was used to skip "
    "malformed rows without stopping the entire parsing process. This makes programs more robust "
    "and suitable for real-world data, which is rarely perfectly clean."
)

add_heading("3.5 Type Hints", level=2)
add_body(
    "Type hints are annotations added to function signatures that indicate the expected types of "
    "arguments and return values. They do not enforce types at runtime but serve as documentation "
    "and enable static analysis tools to catch type errors before execution. Using type hints "
    "consistently makes code easier to understand and maintain."
)
add_page_break()

# 4. CSV PARSING AND PANDAS
add_heading("4. CSV Parsing and Pandas")
add_heading("4.1 CSV as a Data Format", level=2)
add_body(
    "A CSV (Comma-Separated Values) file is a plain text format for representing tabular data. "
    "The first row typically contains headers that name each column, and subsequent rows contain "
    "data values separated by commas. CSV is widely used because it is human-readable, "
    "lightweight, and compatible with almost every data processing tool."
)
add_body(
    "Despite its simplicity, CSV files can contain complications such as missing values, extra "
    "whitespace, inconsistent capitalisation, mixed data types in a single column, and duplicate "
    "rows. Understanding these issues is important before attempting any analysis."
)

add_heading("4.2 Manual Parsing", level=2)
add_body(
    "Manual CSV parsing uses only Python's built-in file I/O capabilities. The file is opened "
    "and read line by line. The header line is split on commas to extract column names. Each "
    "subsequent line is split on commas and zipped with the headers to produce a dictionary. "
    "Type conversion is performed explicitly: score values are cast to integers, and submitted "
    "values are normalised to a consistent form."
)
add_body(
    "The manual approach builds a clear mental model of what a CSV actually is: a text file with "
    "a specific structure. It also makes the programmer aware of edge cases such as empty lines, "
    "trailing whitespace, and rows with the wrong number of columns, all of which must be handled "
    "explicitly."
)

add_heading("4.3 Pandas-Based Loading", level=2)
add_body(
    "Pandas is a Python library that provides the DataFrame, a two-dimensional labelled data "
    "structure similar to a spreadsheet. Loading a CSV with pandas requires a single function "
    "call. Pandas automatically infers column types, handles common edge cases, and provides a "
    "rich set of methods for filtering, grouping, and summarising data."
)
add_body(
    "The key difference between manual parsing and pandas is the level of abstraction. Manual "
    "parsing makes every step explicit, while pandas hides the complexity. For production work, "
    "pandas is far more efficient and less error-prone. For learning, manual parsing reveals "
    "what pandas is doing internally."
)

add_heading("4.4 Summary Generation", level=2)
add_body(
    "After loading the data, both approaches calculated the same set of summary statistics: "
    "total students, number who submitted, number who did not, average score, highest scorer, "
    "lowest scorer among those who submitted, domain-wise averages, and students scoring below "
    "a threshold. Comparing the outputs of both methods confirmed that the implementations were "
    "correct and equivalent."
)
add_page_break()

# 5. DATA CLEANING
add_heading("5. Data Cleaning")
add_heading("5.1 Why Data Cleaning is Necessary", level=2)
add_body(
    "Real-world datasets are almost never clean when first collected. Values may be missing, "
    "inconsistently formatted, duplicated, or outright incorrect. If uncleaned data is used "
    "directly for analysis or machine learning, the results will be unreliable. Data cleaning "
    "is the process of identifying and correcting these issues in a systematic and documented way."
)

add_heading("5.2 Common Data Issues", level=2)
add_body(
    "Duplicate rows occur when the same record is entered more than once. They inflate counts "
    "and skew averages. Removing duplicates is usually the first cleaning step. Inconsistent "
    "categorical values occur when the same category is written in multiple ways, for example "
    "ml, ML, and MACHINE LEARNING all referring to the same domain. These must be standardised "
    "to a single canonical form before grouping or filtering."
)
add_body(
    "Wrong data types occur when a numeric column contains text values such as 'nine' instead "
    "of 9, or when a percentage column contains strings like '92%' instead of the number 92. "
    "These must be converted to the appropriate numeric type before any calculation. Unit "
    "inconsistencies occur when values in the same column use different units, for example "
    "height recorded as both 170 cm and 1.62 m. All values must be converted to a single unit "
    "before comparison or analysis."
)

add_heading("5.3 Handling Missing Values", level=2)
add_body(
    "Missing values must not be ignored silently or removed without justification. Each missing "
    "value requires a decision: impute it with a reasonable value such as the column mean or "
    "median, mark it as unknown, or remove the row if the missing field is critical and cannot "
    "be inferred. The choice and the reasoning must be documented clearly in the cleaning report "
    "so that any reviewer can understand what was done and why."
)

add_heading("5.4 Handling Invalid Values", level=2)
add_body(
    "Invalid values are values that are present but make no sense in context. An attendance "
    "percentage of -10 or 105 is technically a number but is logically impossible. Such values "
    "must either be corrected if the correct value can be inferred, or removed with justification. "
    "Silently keeping invalid values would corrupt any downstream analysis."
)

add_heading("5.5 Validation After Cleaning", level=2)
add_body(
    "After cleaning, the dataset must be validated to confirm that all issues have been resolved. "
    "Validation checks include confirming that no duplicate student IDs exist, that attendance "
    "values fall within the valid range of 0 to 100, that all numeric columns contain only "
    "numbers, that categorical columns contain only the expected values, and that no critical "
    "columns have missing values. Validation transforms cleaning from a best-effort process into "
    "a verifiable one."
)
add_page_break()

# 6. DATA VISUALIZATION
add_heading("6. Data Visualization")
add_heading("6.1 Why Visualization Matters", level=2)
add_body(
    "Numerical summaries describe data but do not always reveal its shape, distribution, or "
    "patterns. Visualization makes patterns immediately visible. A bar chart showing domain-wise "
    "average scores communicates at a glance which domain performed best. A scatter plot showing "
    "the relationship between attendance and score reveals whether students who attend more tend "
    "to score higher. These insights are difficult to extract from a table of numbers alone."
)

add_heading("6.2 The Three Plots", level=2)
add_body(
    "The first plot is a bar chart showing the average score for each domain. This allows a "
    "quick comparison of performance across ML, Web, Electronics, and Mechanical domains. The "
    "x-axis shows domain names and the y-axis shows average score. A clear title and axis labels "
    "ensure the chart is self-explanatory."
)
add_body(
    "The second plot is a scatter plot showing attendance percentage on the x-axis and score on "
    "the y-axis. Each point represents one student. This plot helps identify whether higher "
    "attendance is associated with higher scores, or whether the relationship is weak or absent. "
    "Scatter plots are the appropriate choice when exploring the relationship between two "
    "continuous variables."
)
add_body(
    "The third plot is a bar chart showing the count of students who submitted versus those who "
    "did not. This gives an immediate sense of participation rate. A simple two-bar chart with "
    "clear labels is sufficient for this purpose."
)

add_heading("6.3 Matplotlib Best Practices", level=2)
add_body(
    "Every plot must have a descriptive title so that a reader can understand what it shows "
    "without additional context. Both axes must be labelled with the variable name and unit "
    "where applicable. Plots must be saved using savefig() rather than displayed interactively, "
    "because the code must run in non-interactive environments. Using tight_layout() prevents "
    "labels from being cut off at the edges of the figure. These practices ensure that the "
    "output plots are professional and usable outside the development environment."
)
add_page_break()

# 7. TASK SUMMARY TABLE
add_heading("7. Task Summary")
add_body("The table below summarises each task and the primary concept learned from it.")

table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

headers = ["Task", "Title", "Main Concept Learned"]
widths = [Inches(0.8), Inches(1.6), Inches(4.0)]
rows_data = [
    ("Task 1", "Environment Setup",   "Git, GitHub, virtual environments, Linux command-line basics, and repository hygiene."),
    ("Task 2", "Python Fundamentals", "Functions, lists, dictionaries, file I/O, exception handling, and type hints in Python."),
    ("Task 3", "CSV Parsing",         "Manual CSV parsing with file I/O versus pandas loading; type conversion and summary generation."),
    ("Task 4", "Data Cleaning",       "Handling missing values, duplicates, type errors, unit inconsistencies, and post-cleaning validation."),
    ("Task 5", "Visualization",       "Generating and saving bar charts and scatter plots using matplotlib with proper titles and axis labels."),
]

for i, hdr in enumerate(headers):
    c = table.rows[0].cells[i]
    c.width = widths[i]
    set_cell_bg(c, "000000")
    cell_text(c, hdr, bold=True, size=11, color="FFFFFF")

for r_idx, row_data in enumerate(rows_data):
    for c_idx, val in enumerate(row_data):
        c = table.rows[r_idx + 1].cells[c_idx]
        c.width = widths[c_idx]
        set_cell_bg(c, "F2F2F2" if r_idx % 2 == 0 else "FFFFFF")
        cell_text(c, val, size=10)

add_page_break()

# 8. CONCLUSION
add_heading("8. Conclusion")
add_body(
    "The five tasks completed during this taskphase represent a complete, end-to-end data "
    "processing pipeline. Beginning with environment setup and version control, progressing "
    "through Python scripting and CSV parsing, tackling the challenge of messy real-world data, "
    "and finally producing visual summaries, each task built directly on the knowledge of the "
    "previous one."
)
add_body(
    "The most important insight from this work is that data quality determines the quality of "
    "any analysis. No visualization or model built on uncleaned data can be trusted. Systematic "
    "cleaning, with clear rules and documented decisions, is not optional — it is the foundation "
    "of reliable data work."
)
add_body(
    "Version control with Git ensures that the entire history of the project is preserved and "
    "that the work is reproducible. Virtual environments and requirements.txt ensure that anyone "
    "can recreate the same development environment. Together, these practices reflect the "
    "standards expected in professional software and machine learning projects."
)
add_page_break()

# 9. REFERENCES
add_heading("9. References")
refs = [
    "1. Python Software Foundation. Python 3 Documentation. https://docs.python.org/3/",
    "2. pandas Development Team. pandas User Guide. https://pandas.pydata.org/docs/",
    "3. Hunter, J. D. Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 2007.",
    "4. Chacon, S. and Straub, B. Pro Git Book. https://git-scm.com/book/en/v2",
    "5. GitHub Docs. https://docs.github.com/",
    "6. Python Packaging Authority. pip and virtualenv documentation. https://pip.pypa.io/",
]
for ref in refs:
    add_body(ref)

# SAVE
output_path = "task_6/report/Software_ML_Taskphase_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")